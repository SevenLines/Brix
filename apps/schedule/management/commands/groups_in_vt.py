# Контингенты 1-3 курсов в компьютерных классах
from itertools import groupby
from pprint import pprint

from django.core.management import BaseCommand
from django.db.models import F, Q
from django.db.models.functions import Trim
from docx import Document

from apps.schedule.models import Kontgrp, Ownres, Raspis, Auditory, AudList, AudSps


class Command(BaseCommand):
    help = 'Closes the specified poll for voting'

    def add_arguments(self, parser):
        pass

    def handle(self, *args, **options):
        groups = Kontgrp.objects.filter(
            kont__kurs__in=(1, 2, 3),
            kont__id__in=Ownres.objects.filter(
                ownerid__in=(
                    9, 18, 21, 17, 16,
                    24, 13, 7, 12, 10, 41,
                    8, 20, 11, 14, 43, 28,
                    36, 35, 34, 40, 32, 26,
                    31, 29, 42, 27, 38, 30, 33,
                )
            ).values("objid"),
        ).order_by("title")

        raspis = Raspis.objects.annotate(
            kont_id=F("raspnagr__kontid"),
            prep=Trim(F("raspnagr__prep__short_name")),
        ).filter(
            Q(aud__in=Auditory.objects.filter(
                Q(korp=19) | Q(id__in=[
                    489, 413, 367, 369, 64, 634
                ])
            ).values("id")) | Q(
                aud__in=AudList.objects.filter(aud__in=[
                    489, 413, 367, 369, 64, 634
                ]).values("auds")
            ),
            raspnagr__kontid__in=groups.values("id"),
        ).order_by("raspnagr__kontid__title", "everyweek", "day", "para")

        raspis = {kont_id: list(items) for kont_id, items in groupby(raspis, lambda x: x.kont_id)}

        document = Document()

        auditories = {i.id: i for i in Auditory.objects.all()}
        auditories_list = {i.id: i for i in AudSps.objects.all()}

        for group in groups:
            raspis_items = raspis.get(group.id, [])
            if raspis_items:
                p = document.add_paragraph()
                r = p.add_run(group.title)
                r.font.bold = True

                for item in raspis_items:
                    every_week = "еженедельная" if item.everyweek == 2 else "нечетная" if item.day <= 7 else "четная"
                    day = {
                        1: "понедельник",
                        2: "вторник",
                        3: "среда",
                        4: "четверг",
                        5: "пятница",
                        6: "суббота",
                    }.get(item.day % 7)

                    para = "{}-ая пара".format(item.para)

                    aud_title = auditories.get(item.aud, auditories_list.get(item.aud)).obozn

                    p.add_run("\n{}".format(" ".join([
                        str(aud_title.strip()),
                        str(every_week),
                        str(day),
                        str(para),
                        str(item.prep),
                    ])))
            else:
                print(group.title)

        # document.save("konts.docx")
